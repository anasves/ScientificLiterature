__author__ = 'anastasia'
from Bio import Entrez
from Bio import Medline
from docx import Document

document = Document()

document.add_heading('Document Title', 0)

MAX_COUNT = 100
#TERM = 'Vassily Hatzimanikatis'
#TERM = 'Frances H. Arnold'
TERM = 'Chimeric Enzymes'

print('Getting {0} publications containing {1}...'.format(MAX_COUNT, TERM))
Entrez.email = 'A.N.Other@example.com'
h = Entrez.esearch(db='pubmed', retmax=MAX_COUNT, term=TERM)
result = Entrez.read(h)
print('Total number of publications containing {0}: {1}'.format(TERM, result['Count']))
# Total number of publications containing Vassily Hatzimanikatis: 77
ids = result['IdList']

#97 publications in pubmed containing Hatzimanikatis

h = Entrez.efetch(db='pubmed', id=ids, rettype='medline', retmode='text')

records = Medline.parse(h)

# Record
# {'PMID': '17098800',
# 'OWN': 'NLM',
# 'STAT': 'MEDLINE',
# 'DCOM': '20070306',
# 'LR': '20181113',
# 'IS': '0006-3495 (Print) 0006-3495 (Linking)',
#  'VI': '92',
#  'IP': '3',
# 'DP': '2007 Feb 1',
#  'TI': 'A model for protein translation: polysome self-organization leads to maximum protein synthesis rates.',
# 'PG': '717-30',
# 'AB': 'The genetic information in DNA is transcribed to mRNA and then translated to proteins, which form the building blocks of life. Translation, or protein synthesis, is hence a central cellular process. We have developed a gene-sequence-specific mechanistic model for the translation machinery, which accounts for all the elementary steps of the translation mechanism. We performed a sensitivity analysis to determine the effects of kinetic parameters and concentrations of the translational components on protein synthesis rate. Utilizing our mathematical framework and sensitivity analysis, we investigated the translational kinetic properties of a single mRNA species in Escherichia coli. We propose that translation rate at a given polysome size depends on the complex interplay between ribosomal occupancy of elongation phase intermediate states and ribosome distributions with respect to codon position along the length of the mRNA, and this interplay leads to polysome self-organization that drives translation rate to maximum levels.',
# 'FAU': ['Zouridis, Hermioni', 'Hatzimanikatis, Vassily'],
# 'AU': ['Zouridis H', 'Hatzimanikatis V'],
# 'AD': 'Department of Chemical and Biological Engineering, McCormick School of Engineering and Applied Sciences, Northwestern University, Evanston, Illinois, USA.',
# 'LA': ['eng'],
# 'PT': ['Journal Article', "Research Support, Non-U.S. Gov't", "Research Support, U.S. Gov't, Non-P.H.S."],
# 'DEP': '20061110',
# 'PL': 'United States',
# 'TA': 'Biophys J',
# 'JT': 'Biophysical journal',
# 'JID': '0370626',
# 'RN': ['0 (Escherichia coli Proteins)', '0 (Transcriptional Elongation Factors)', '9014-25-9 (RNA, Transfer)'],
# 'SB': 'IM',
# 'MH': ['Computer Simulation', 'Escherichia coli/physiology', 'Escherichia coli Proteins/*biosynthesis', '*Models, Biological', 'Polyribosomes/*metabolism', 'Protein Biosynthesis/*physiology', 'RNA, Transfer/*metabolism', 'Transcriptional Elongation Factors/*metabolism'],
# 'PMC': 'PMC1779991',
# 'EDAT': '2006/11/14 09:00',
# 'MHDA': '2007/03/07 09:00',
# 'CRDT': ['2006/11/14 09:00'],
# 'PHST': ['2006/11/14 09:00 [pubmed]', '2007/03/07 09:00 [medline]', '2006/11/14 09:00 [entrez]'],
# 'AID': ['S0006-3495(07)70883-4 [pii]', '10.1529/biophysj.106.087825 [doi]'],
# 'PST': 'ppublish',
# 'SO': 'Biophys J. 2007 Feb 1;92(3):717-30. doi: 10.1529/biophysj.106.087825. Epub 2006 Nov 10.'}

affiliations = []
for record in records:
    if 1 == 1:
    #affiliation = record.get('AD')
    #if affiliation and 'Laboratory of Computational Systems Biotechnology' in affiliation:
        year = record.get('EDAT')
        journal = record.get('JT')
        title = record.get('TI')
        abstract = record.get('AB')
        authors = record.get('FAU')
        document.add_heading('{}'.format(title), level=1)
        document.add_paragraph(year[0:4])
        document.add_paragraph(journal)
        for author in authors:
            p = document.add_paragraph(author)
        p = document.add_paragraph(abstract)
        document.add_page_break()
print('\n'.join(affiliations))

document.save('Chimeric enzymes abstracts pubmed.docx')